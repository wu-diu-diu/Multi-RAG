from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Union
import os
import fitz  # PyMuPDF
import docx
import openpyxl
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

class DocumentProcessor(ABC):
    """文档处理基类"""
    
    def __init__(self, config: Dict):
        self.config = config
        self.supported_formats = config['document_processor']['supported_formats']
        self.ocr_config = config['document_processor']['ocr']    
    @abstractmethod
    def process(self, file_path: str) -> Dict:
        """处理文档并返回结构化数据"""
        pass
    
    def extract_text(self, file_path: str) -> str:
        """提取文档中的文本"""
        pass
    
    def extract_images(self, file_path: str) -> List[Image.Image]:
        """提取文档中的图片"""
        pass
    
    def extract_tables(self, file_path: str) -> List[Dict]:
        """提取文档中的表格"""
        pass
    
    def perform_ocr(self, image: Image.Image) -> str:
        """对图片进行OCR识别"""
        return pytesseract.image_to_string(
            image,
            lang=self.ocr_config['language'],
            config=f'--dpi {self.ocr_config["dpi"]}'
        )

class PDFProcessor(DocumentProcessor):
    """PDF文档处理器"""
    
    def process(self, file_path: str) -> Dict:
        result = {
            'text': self.extract_text(file_path),
            'images': self.extract_images(file_path),
            'tables': self.extract_tables(file_path)
        }
        return result
    
    def extract_text(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    
    def extract_images(self, file_path: str) -> List[Image.Image]:
        images = convert_from_path(file_path)
        return images
    
    def extract_tables(self, file_path: str) -> List[Dict]:
        # 使用PyMuPDF提取表格
        doc = fitz.open(file_path)
        tables = []
        for page in doc:
            # 这里需要实现表格提取逻辑
            pass
        return tables

class WordProcessor(DocumentProcessor):
    """Word文档处理器"""
    
    def process(self, file_path: str) -> Dict:
        result = {
            'text': self.extract_text(file_path),
            'images': self.extract_images(file_path),
            'tables': self.extract_tables(file_path)
        }
        return result
    
    def extract_text(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def extract_images(self, file_path: str) -> List[Image.Image]:
        # 实现Word文档中的图片提取
        pass
    
    def extract_tables(self, file_path: str) -> List[Dict]:
        doc = docx.Document(file_path)
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)
        return tables

class ExcelProcessor(DocumentProcessor):
    """Excel文档处理器"""
    
    def process(self, file_path: str) -> Dict:
        result = {
            'text': self.extract_text(file_path),
            'tables': self.extract_tables(file_path)
        }
        return result
    
    def extract_text(self, file_path: str) -> str:
        # Excel通常不包含普通文本
        return ""
    
    def extract_tables(self, file_path: str) -> List[Dict]:
        wb = openpyxl.load_workbook(file_path)
        tables = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            table_data = []
            for row in ws.rows:
                table_data.append([cell.value for cell in row])
            tables.append({
                'sheet_name': sheet,
                'data': table_data
            })
        return tables 